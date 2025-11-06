from __future__ import annotations
from pathlib import Path
from typing import List
from .base import BaseParser, ParseResult, Table


class DocxParser(BaseParser):
    supported_suffixes = [".docx"]

    def parse(self, path: Path) -> ParseResult:
        from docx import Document  # python-docx
        result = ParseResult(file_path=path)
        doc = Document(str(path))

        paragraphs: List[str] = []
        tables: List[Table] = []

        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                paragraphs.append(text)

        for idx, t in enumerate(doc.tables, start=1):
            rows = []
            for row in t.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            tables.append(Table(source=f"docx_table_{idx}", rows=rows))

        result.pages = []  # not applicable
        result.text = "\n".join(paragraphs).strip()
        result.tables = tables
        # basic metadata
        core_props = doc.core_properties
        meta = {
            "author": core_props.author,
            "title": core_props.title,
            "created": str(core_props.created) if core_props.created else None,
        }
        result.metadata.update({k: v for k, v in meta.items() if v})
        return result
